"""
Функции для генерации выходного файла с оформленным списком использованных источников.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


class Renderer:
    """
    Создание выходного файла – Word.
    """

    def __init__(self, rows: tuple[str, ...]):
        self.rows = rows

    def set_style(self, document: Any) -> str | None:
        """
        Метод получения стилей для выходного файла.
        """
        raise NotImplementedError

    def render(self, path: Path | str) -> None:
        """
        Метод генерации Word-файла со списком использованных источников.

        :param Path | str path: Путь для сохранения выходного файла.
        """

        document = Document()

        # стилизация заголовка
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        runner = paragraph.add_run("Список использованной литературы")
        runner.bold = True

        # стилизация текста
        style_normal = document.styles["Normal"]
        style_normal.font.name = "Times New Roman"
        style_normal.font.size = Pt(12)
        style_normal.paragraph_format.line_spacing = 1.5
        style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for row in self.rows:
            # добавление источника
            document.add_paragraph(row, style="List Number")

        # сохранение файла Word
        document.save(path)


class GOSTRenderer(Renderer):
    def set_style(self, document: Any) -> str | None:
        """
        Метод задания стиля для оформления текста.
        """
        style_normal = document.styles["Normal"]

        style_normal.font.name = "Times New Roman"
        style_normal.font.size = Pt(12)
        style_normal.paragraph_format.line_spacing = 1.5
        style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        return "List Number"


class APARenderer(Renderer):
    def set_style(self, document: Any) -> str | None:
        """
        Метод задания стиля для оформления текста.
        """
        style_normal = document.styles["Normal"]

        style_normal.font.name = "Times New Roman"
        style_normal.font.size = Pt(12)
        style_normal.paragraph_format.line_spacing = 1.5
        style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        style_normal.paragraph_format.first_line_indent = Cm(-1.5)
        return None
